# Databricks notebook source
# MAGIC %md
# MAGIC **Lib**

# COMMAND ----------

! pip3 install langchain unstructured[all-docs] pydantic lxml poppler-utils unstructured[azure] PyMuPDF pdf2docx langchain_openai weasyprint pdfkit Spire.Doc pillow docx 

# sudo apt-get install poppler-utils tesseract-ocr

# cd /
# cp /Workspace/Users/jfeng51@its.jnj.com/Transformed/builder.py databricks/python/lib/python3.10/site-packages/google/protobuf/internal/

# COMMAND ----------

! pip install --upgrade pip

# COMMAND ----------

# MAGIC %md
# MAGIC **PDF_2_IMAGE**

# COMMAND ----------

# MAGIC %md
# MAGIC Make sure to run `sudo apt install poppler-utils` in terminal.

# COMMAND ----------

from pdf2image import convert_from_path
import os
 
# file_name = "942249.pdf"
# file_name = "1315975.pdf"
# file_name = "887569.pdf"
# file_name = "889986.pdf"
# file_name = "1482067.pdf"
file_name = "Ustekinumab - Occurrence of Alopecia in Adult Patients with CD or UC Treated with Ustekinumab.pdf"
image_path = "/Workspace/Users/dsheng3@its.jnj.com/JP FAQ/Transformed/Global_Image/" + file_name.split('.')[0] + "/"
pdf_path = "/Workspace/Users/jfeng51@its.jnj.com/Transformed/PDF/"

pdf_images_list = []
 
# 将PDF文件的每一页转换为图像
images = convert_from_path(pdf_path + file_name)

if not os.path.exists(image_path):
    os.makedirs(image_path)
 
# 可以遍历images列表来保存或处理每一张图像
for i, image in enumerate(images):
    # 图像保存路径
    image_name = f'page_{i + 1}.png'
    # 保存图像
    image.save(image_path + image_name, 'PNG')
    print(image_path, image_name)
    pdf_images_list.append(image_path + image_name)

# COMMAND ----------

displayHTML("<img src ='/Global_Image/1315975/page_1.png'>")

# COMMAND ----------

!pip install streamlit

# COMMAND ----------

import streamlit as st

image_url = "https://adb-6198580363974626.6.azuredatabricks.net/?o=6198580363974626#files/4071630213144838.png"
st.image(image_url)

# COMMAND ----------

from IPython.display import Image 
from IPython.core.display import HTML 
Image(url= "Global_Image/1315975/page_1.png")

# COMMAND ----------

import matplotlib.pyplot as plt
import matplotlib.image as mpimg

# 指定图像的完整路径
image_path = '/Workspace/Users/dsheng3@its.jnj.com/github/JP_FAQ/Global_Image/1315975/page_1.png'

# 使用 matplotlib 加载和显示图像
img = mpimg.imread(image_path)
plt.imshow(img)
plt.axis('off')  # 关闭坐标轴
plt.show()


# COMMAND ----------

# 指定图像在工作区中的路径
image_path = '/Workspace/Users/dsheng3@its.jnj.com/github/JP_FAQ/Global_Image/1315975/page_1.png'

# 生成 HTML 代码来显示图像
displayHTML(f'<img src="{image_path}">')


# COMMAND ----------

# MAGIC %md
# MAGIC **UTILS**

# COMMAND ----------

# TODO
# 按顺序将, image,table之前,之后以及相关的text保存并填充
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
import re

# 返回字符串位置
def find_str_positions(str, parse_str):
    positions = []
    index = str.find(parse_str)
    while index!= -1:
        positions.append(index)
        index = str.find(parse_str, index + 1)
    return positions

# 找表格||之间的'\n', 即找table行数
def count_newlines_between_pipes(str):
    first_pipe = str.find('|')
    last_pipe = str.rfind('|')
    if first_pipe == -1 or last_pipe == -1:
        return 0
    count = str[first_pipe:last_pipe].count('\n')
    return count

# 找表格|\n|之间的'|', 即找table列数
def count_segement_between_pipes(str):
    first_pipe = str.find("|")
    last_pipe = str.find("|\n")
    if first_pipe == -1 or last_pipe == -1:
        return 0
    count = str[first_pipe:last_pipe].count('|')
    return count

# 正则表达判断(是否为字符)
def is_all_letters_regex(string):
    pattern = re.compile('^[a-zA-Z0-9]+$')
    return bool(pattern.match(string))

# 正则表达式判断(是否为非字符)
def is_only_non_letters_regex(string):
    pattern = re.compile('^[^a-zA-Z0-9]+$')
    return bool(pattern.match(string))

# 正则表达式判断(提取两个匹配的字符串中间的字符串)
def extract_between_strings(text, start_str, end_str):
    pattern = f"{re.escape(start_str)}(.*?){re.escape(end_str)}"
    match = re.search(pattern, text, re.DOTALL)
    if match:
        return match.group(1)
    else:
        return ""
    
# 正则表达式匹配'-'
def match_only_hyphens(s):
    pattern = r'^---$'
    if re.match(pattern, s):
        return True
    else:
        return False

# 设置单元格边框
def set_cell_border(cell: _Cell, **kwargs):
    """
    设置单元格边框

    参数：
    - cell (_Cell)：要设置边框的单元格
    - kwargs：边框属性，例如 top、bottom、left、right 等

    用法：
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#ff0000", "space": "0"},
        bottom={"sz": 12, "color": "#00ff00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 检查是否存在标签，如果不存在，则创建一个
    tcBorders = tcPr.first_child_found_in("w:tcborders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcborders')
        tcPr.append(tcBorders)

    # 遍历所有可用的标签
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # 检查是否存在标签，如果不存在，则创建一个
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # 设置属性
            for key in ("sz", "val", "color", "space", "shadow"):
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))   

def create_table(r, c):
    table = doc.add_table(rows=r, cols=c)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  #设置新表格的对齐方式
    # Set New Table Boarder
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 
                            top={"sz": 2, "val": "single", "color": "000000"}, 
                            bottom={"sz": 2, "val": "single", "color": "000000"},
                            left={"sz": 2, "val": "single", "color": "000000"},
                            right={"sz": 2, "val": "single", "color": "000000"})
    return table

def check_string_regex(s):
    pattern = r'^[|-]+$'
    if re.match(pattern, s):
        return True
    else:
        return False

# 找到所有的 key_word 出现的位置, 提取keyword之间字符串
def extract_keyword_strings(text, key_word):
    extract_contet = []
    table_indices = [m.start() for m in re.finditer(key_word, text)]
    if len(table_indices) < 2:
        return "没有足够的 TABLE 来截取"
    for key_word_index in range(len(table_indices) - 1):
        star_index = table_indices[key_word_index] + len(key_word)
        end_index = table_indices[key_word_index + 1]
        extract_contet.append(text[star_index:end_index].strip())
    return extract_contet

# COMMAND ----------

# MAGIC %md
# MAGIC **GPT Process**

# COMMAND ----------

# TODO
# 4o会将Image和Table搞混, 需要一个办法强制输出正确格式, 或删除Image相关的Prompt, 处理4o自然输出

import base64
import requests
from mimetypes import guess_type
from pdf2image import convert_from_path
from docx.shared import Inches
import os
 
# file_name = "Ustekinumab - Occurrence of Alopecia in Adult Patients with CD or UC Treated with Ustekinumab.pdf"
# file_name = "Ustekinumab - The Use of Vaccines in Infants Exposed In Utero to Ustekinumab.pdf"
# file_name = "Ustekinumab - Use in Adult Patients with Comorbid Liver Cirrhosis.pdf"
# file_name = "Ustekinumab - Use in Adult Patients with Crohn's Disease or Ulcerative Colitis and Comorbid PSC.pdf"
# file_name = "Ustekinumab - Occurrence of Cytomegalovirus in Adult Patients Receiving Ustekinumab.pdf"
file_name = "Ustekinumab - Occurrence of Skin-Related Adverse Events in Crohn's Disease and Ulcerative Colitis.pdf"

image_path = "/Workspace/Users/jfeng51@its.jnj.com/Transformed/Global_Image/" + file_name.split('.')[0] + "/"
pdf_path = "/Workspace/Users/jfeng51@its.jnj.com/Transformed/PDF/"

pdf_images_list = []
 
# 将PDF文件的每一页转换为图像
images = convert_from_path(pdf_path + file_name)

if not os.path.exists(image_path):
    os.makedirs(image_path)
 
# 可以遍历images列表来保存或处理每一张图像
for i, image in enumerate(images):
    # 图像保存路径
    image_name = f'page_{i + 1}.png'
    # 保存图像
    image.save(image_path + image_name, 'PNG')
    pdf_images_list.append(image_path + image_name)


# Encode the image to base64
sImageData_list = []
for image_path in pdf_images_list:
    sImageData = base64.b64encode(open(image_path, 'rb').read()).decode('utf-8')
    sImageData_list.append(sImageData)

# API Initialize
sEndpoint='https://azr-hjg-apac-ds-cae-openai-3.openai.azure.com/'
sKey='d3460cd8602e4a8d8420ba021b886cd3'
sDeployment='gpt-4o'

# dData is copied from the tutorial, Temelate
dData = {
    "messages": [
        {
            "role": "system",
            "content": "You are a professional Japanese translators for STL related treatments, your translated article is very fluent and easy to understand"
        },
        {
            "role": "user",
            "content": [
                {   
                    "type": "text",

                    "text": 
                    "## Role: Translator of specialized medical knowledge"
                    
                    "## Background: The following image shows a medical report on the drug Ustekinumab, which contains a title (Ustekinumab + subject), as well as a number of paragraph headings such as Review, Clinical Data, Literature Index, References, etc."
                    
                    "## Purpose: Please translate the medical report on Ustekinumab (drug name) shown in the image below into professional scientific Japanese, according to the following rules for layout and translation output"

                    "## Translations rules:"
                    "- All contents need to be translated into Japanese"
                    "- Translations should be done in the original language to make them sound" 
                    "- Avoid honorific expressions in Japanese translations, use objective statements"
                    "- Streamline the translation"
                    "- Delete unnecessary repetitive subjects."
                    
                    "## Layout rules:"
                    "- Generate Word with the same layout as in the original image"
                    "- Output is formatted as, ```markdown line breaks Output content line breaks ```"
                    "- Left-aligned main headings, formatted as drug name line feed title, drug name in original, subject needs translation"
                    "- Split the output Word into four paragraphs, the four paragraphs are titled 'Offer, Clinical Data, Literature Search, Reference', Offer is the SUMMARY or COMPANY CORE DATA SHEET paragraph in the original text, Clinical Data is the CLINICAL DATA or TECHNICAL MEMORANDUM paragraph in the original text, Literature Search is the Literature Search in the original text, Literature Search is the Literature Search in the original text. Search is the LITERATURE SEARCH paragraph in the original text, and Reference is the REFERENCES paragraph in the original text."
                    "- If there is no Literature Search content, output '(this セクションは画像に含まれていません)'"
                    "- If there is no reference content, output '(このセクションは画像に含まれていません)'"
                    # “- Literature search and references do not need to be translated, use the original text.”
                    "- Paragraph headings need to be translated into Japanese."
                    "- Paragraph headings are left-justified"
                    "- Delete the footer"
                    "- The '-' and 'o' symbols in the original text are retained in the output."
                    "- If there is a table, add the word 'TABLE' at the beginning and at the end of the output table, don't forget the table description and the abbreviations, in the following format: 'TABLE' line feed table description line feed table (Important: the contents of the table should be translated into Japanese) line feed abbreviations (Important: all the contents of the Abbreviations) 'TABLE' (Important: the end of the ' TABLE' (important: end of ' TABLE' must be output), Cell is split by '|', Cell with no content is displayed by |---|, Cell with empty content is displayed by |empty|, text in Cell needs to be translated into Japanese and output."
                    "- If there is no table, do not output 'TABLE'"
                    # “- If there is an image, add the word ‘IMAGE’ at the beginning and at the end of the output image, don't forget the image description and the abbreviation, in the format IMAGE line feed 図:圖片描述 line feed  略语:Content of Abbreviations: line feed IMAGE”
                    # “- if there is no image, do not output ‘IMAGE’”
                    "- Tables are formatted the same way as in the original image, content is converted to Japanese"
                    "- Cells of the table are spaced with |"
                    "- Cells in tables are spaced with '|'"
                    "- Line breaks in tables are indicated by '-'"

                    "## Vocabulary rules:"
                    "- English:Guttate Psoriasis, Japanese:乾癬"
                    "- English: chronic plaque PsO, Japanese: 尋常性乾癬."
                    "- English:Ustekinumab, Japanese:ステラーラ"
                },
            ]
        },
        # {
        #     "role": "assistant",
        #     "content": "'The table has a total of"
        # }
    ],
    "max_tokens": 4096, #max_token
    "stream": False,
    "temperature": 0,
}

# Import PDF Image
for image_data in sImageData_list:
    url_data = {"type": "image_url", 
                "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
    dData["messages"][1]["content"].append(url_data)

# Make the API request
response = requests.post(
    f'{sEndpoint}openai/deployments/{sDeployment}/chat/completions?api-version=2024-02-01',
    headers={'api-key': sKey, 'Content-Type': 'application/json'},
    json=dData
)

# Print the response
print(response.json()['choices'][0]['message']['content'])

# COMMAND ----------

response.json()

# COMMAND ----------

# MAGIC %md
# MAGIC **HANDLE DATA**

# COMMAND ----------

"""
Multi Table
"""
def parse_table(clinical_text, cell):
    clinical_text_list = clinical_text.split("TABLE")
    del(clinical_text_list[0])
    del(clinical_text_list[-1])
    table_list = extract_keyword_strings(clinical_text, "TABLE")
    for table_index in range(len(table_list)):
        if table_index % 2 == 0:
            """ 获取'TABLE'间内容 """
            table_text = table_list[table_index]

            """ TABLE描述和略语 """
            table_describle_text = table_text.split('|')[0].strip()
            # print(table_describle_text)
            abbreviation_text = table_text[table_text.rfind('|'):].replace("|", "").strip()
            # print(abbreviation_text)

            """ CREATE TABLE """
            # Create New Table
            new_table_row = count_newlines_between_pipes(table_text.strip())
            print("new_talbe_row: " + str(new_table_row))
            new_table_col = count_segement_between_pipes(table_text.strip())
            print("new_talbe_col: " + str(new_table_col))
            new_table = create_table(new_table_row, new_table_col)
            new_table.alignment = WD_TABLE_ALIGNMENT.CENTER  #设置新表格的对齐方式

            """ 分解TABLE内容 """
            # 方案一 '|'切割
            # table_text_list = table_text[table_text.find("|"): table_text.rfind("|")].strip().split("|")
            # table_text_list = [item.strip() for item in table_text_list if item.strip() not in ["\n", '']]

            # 方案二 '\n'切割
            table_content = (table_text[table_text.find("|"): table_text.rfind("|")] + '|').strip()
            row_list_content = table_content.split("\n")
            row_list_content = [item for item in row_list_content if not check_string_regex(item)]
            # print(row_list_content)

            """ 根据PDF TABLE内容填充到NewTable """
            for row in range(len(new_table.rows)):
                content_list = []
                if row_list_content[row].count('|') != new_table_col + 1:
                    content_list = row_list_content[row].split('|')
                    del(content_list[0])
                    del(content_list[-1])
                    for i in range(new_table_col - len(content_list)):
                        content_list.append("")
                    print(content_list)
                else:
                    content_list = row_list_content[row].split('|')
                    del(content_list[0])
                    del(content_list[-1])
                    print(content_list)
                for col in range(len(new_table.columns)):
                    cur_text = content_list[col]
                    new_table.cell(row, col).text = content_list[col].replace("-•", "\n•").strip()
                    print(new_table.cell(row, col).text)
                    
            """ 合并段落 """
            cell.add_paragraph(table_describle_text)
            paragraph = cell.add_paragraph()
            paragraph._p.addnext(new_table._tbl)
            cell.add_paragraph(abbreviation_text)
        else:
            cell.add_paragraph(table_list[table_index].strip())
            # print(table_list[table_index])

def parse_data_no_table(main_table, summary_text, clinical_text, literature_text, reference_text):
    for row in main_table.rows:
        if row.cells[0].text == "QA番号":
            row.cells[1].text = "QA-0000001"
        elif row.cells[0].text == "領域/製品":
            row.cells[1].text = "免疫領域/ステラーラ/STL"
        elif row.cells[0].text == "QA種別":
            row.cells[1].text = "手動記入"
        elif row.cells[0].text == "質問":
            row.cells[1].text = "[STL]" + title_name
        elif row.cells[0].text == "回答\n（書式設定あり）":
            row.cells[1].paragraphs[2].text = summary_text.strip()
            row.cells[1].add_paragraph(clinical_text.strip())
            # row.cells[1].paragraphs[4].text = clinical_text.strip()
        elif row.cells[0].text == "引用文献":
            row.cells[1].text = "文献検索:\n" + literature_text + "\n" +"参考:\n" + reference_text
            row.cells[1].text = row.cells[1].text.strip()


def parse_data(main_table, summary_text, clinical_text, literature_text, reference_text):
    for row in main_table.rows:
        if row.cells[0].text == "QA番号":
            row.cells[1].text = "QA-0000001"
        elif row.cells[0].text == "領域/製品":
            row.cells[1].text = "免疫領域/ステラーラ/STL"
        elif row.cells[0].text == "QA種別":
            row.cells[1].text = "手動記入"
        elif row.cells[0].text == "質問":
            row.cells[1].text = "[STL]" + title_name
        elif row.cells[0].text == "回答\n（書式設定あり）":
            """ summary """
            row.cells[1].paragraphs[2].text = summary_text.strip()

            """ clinical """
            before_table_text = clinical_text[:clinical_text.find("TABLE")].strip()
            after_table_text = clinical_text[clinical_text.rfind("TABLE"):].replace("TABLE", "").strip()
            
            # print(before_table_text)
            # print(after_table_text)

            row.cells[1].add_paragraph(before_table_text)
            parse_table(clinical_text, row.cells[1])
            row.cells[1].add_paragraph(after_table_text)

        elif row.cells[0].text == "引用文献":
            row.cells[1].text = "文献検索:\n" + literature_text + "\n" + "参考:\n" + reference_text
            row.cells[1].text = row.cells[1].text.strip()

# COMMAND ----------

# MAGIC %md
# MAGIC **Main Process**

# COMMAND ----------

# 指定 Word 文档路径
doc_path = '/Workspace/Users/jfeng51@its.jnj.com/Transformed/Template/Japanese_Template_1.docx'
save_path = '/Workspace/Users/jfeng51@its.jnj.com/Transformed/Word/' + file_name.split('.')[0] + '/'
if not os.path.exists(save_path):
    os.makedirs(save_path)
save_path = save_path + file_name.split('.')[0] + '.docx'

# Parse Data
text = response.json()['choices'][0]['message']['content'].split("```markdown")[1].split("```")[0].strip()
# print(text)

drag_name = text.split("\n")[0].strip()
drag_name = drag_name.replace("#", "").strip()
# print(drag_name + "\n")

title_name = text.split("\n", 2)[1].strip()
title_name = title_name.replace('#', "").strip()
# print(title_name + "\n")

summary_text = extract_between_strings(text, "要約", "臨床データ")
summary_text = summary_text.replace("#", "").strip()
# print(summary_text + "\n")

clinical_text = extract_between_strings(text, "臨床データ", "文献検索")
clinical_text = clinical_text.replace("#", "").strip()
# print(clinical_text + "\n")

literature_text = extract_between_strings(text, "文献検索", "参考")
literature_text = literature_text.replace("#", "").strip()
# print(literature_text + "\n")

reference_text = text.split("参考", 1)[1].strip()
# print(reference_text + "\n")

doc = Document(doc_path)
main_table = doc.tables[0]
target_cell = main_table.cell(4,1)

if clinical_text.find("TABLE") != -1:
    parse_data(main_table, summary_text, clinical_text, literature_text, reference_text)
else:
    parse_data_no_table(main_table, summary_text, clinical_text, literature_text, reference_text)

# Save
doc.save(save_path)

# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------

!sudo apt install poppler-utils
