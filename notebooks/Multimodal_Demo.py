# Databricks notebook source
! pip3 install langchain unstructured[all-docs] pydantic lxml poppler-utils unstructured[azure] PyMuPDF pdf2docx langchain_openai weasyprint pdfkit Spire.Doc pillow docx 

# sudo apt-get install poppler-utils tesseract-ocr

# cd /
# cp /Workspace/Users/jfeng51@its.jnj.com/Transformed/builder.py databricks/python/lib/python3.10/site-packages/google/protobuf/internal/

# COMMAND ----------

# MAGIC %md
# MAGIC PDF_2_Image

# COMMAND ----------

from pdf2image import convert_from_path
import os
 
# file_name = "942249.pdf"
# file_name = "1315975.pdf"
# file_name = "887569.pdf"
# file_name = "889986.pdf"
# file_name = "1482067.pdf"
file_name = "Ustekinumab - Occurrence of Alopecia in Adult Patients with CD or UC Treated with Ustekinumab.pdf"
image_path = "/Workspace/Users/jfeng51@its.jnj.com/Transformed/Global_Image/" + file_name.split('.')[0] + "/"
pdf_path = "/Workspace/Users/jfeng51@its.jnj.com/Transformed/PDF/"

pdf_images_list = []
 
# 将PDF文件的每一页转换为图像
images = convert_from_path(pdf_path + file_name)

if not os.path.exists(image_path):
    os.makedirs(image_path)
 
# 可以遍历images列表来保存或处理每一张图像
for i, image in enumerate(images):
    # 图像保存路径
    image_name = f'page_{i + 1}.png'
    # 保存图像
    image.save(image_path + image_name, 'PNG')
    pdf_images_list.append(image_path + image_name)

# COMMAND ----------

# MAGIC %md
# MAGIC Generate Result

# COMMAND ----------

import base64
import requests
from mimetypes import guess_type

# Encode the image to base64
sImageData_list = []
for image_path in pdf_images_list:
    sImageData = base64.b64encode(open(image_path, 'rb').read()).decode('utf-8')
    sImageData_list.append(sImageData)

# API Initialize
sEndpoint='https://azr-hjg-apac-ds-cae-openai-3.openai.azure.com/'
sKey='d3460cd8602e4a8d8420ba021b886cd3'
sDeployment='gpt-4o'

# dData is copied from the tutorial, Temelate
dData = {
    "messages": [
        {
            "role": "system",
            "content": "You are a professional Japanese translators for STL related treatments, your translated article is very fluent and easy to understand"
        },
        {
            "role": "user",
            "content": [
                {   
                    "type": "text",

                    #HTML Prompt
                    # "text": 
                    # "Extract the content from the following The first three images"
                    # "1.Content according to the original text."
                    # "2.Layout according to the original text"
                    # "3.Translate the Content to Japanese"
                    # "4.keep original image format, output Japanese HTML format"
                    # "5.HTML Style same as last image, head-title is normal, section-title background-color is #0070c0Color,section title color is white, Left aligned text, Don't Forget main title"

                    #Word Prompt
                    # "text": 
                    # "Extract the content from the following images"
                    # "1.Content according to the original text."
                    # "2.Layout according to the original text"
                    # "3.Translate the Content to Japanese"
                    # "4.keep original image format,output Japanes"
                    # "5.If there is a Table in the Image, delete the Table and insert 'The Table at the end' at Table position"

                    #HTML Prompt
                    # "text": 
                    # "## 角色: 专业医疗相关知识翻译员"
                    # "## 背景: 下列图片是一份医疗报告, 其中包含一个标题(药物+主题), 以及一些段落标题例如综述, 临床数据, 文献索引, 参考等"
                    # "## 目的: 将下列图片按照格式翻译为日文,并转换成HTML格式"
                    # "## 规则:"
                    # "- 按照原始图片中的排版生成HTML"
                    # "- 按照原文进行语句通顺的翻译, 不要生成自己的理解"
                    # "- 标题居中显示, 药物名在第一行, 主题在第二行, 药物名用原文, 主题需要翻译为日文"
                    # "- 文献检索内容不需要翻译用原文"
                    # "- 段落标题需要翻译"
                    # "- 段落标题左对齐"
                    # "- 删除掉页脚"
                    # # "- '•'是一个ListItem起始的标志, 'o'是上一个ListItem的ListItem"
                    # "- HTML Style:"
                    # "-- <1> body {font-family: Arial}"
                    # "-- <2> .h1-style {text-align:center; font-weight:bold; font-size:30: padding:30px}"
                    # "-- <3> .section-title {font-weight: bold; color:white; background-color:#0056a0; padding:10px}"
                    # "-- <4> .summary, .clinical data, .literature-search, .references {margin-bottom: 20px;}"
                    # # "-- <5> .p {font_size:5, padding: 5px;}"

                    #Word Prompt
                    "text": 
                    "## 角色: 专业医疗知识翻译员"
                    "## 背景: 下列图片是一份关于Ustekinumab药物的医疗报告, 其中包含一个标题(Ustekinumab + 主题), 以及一些段落标题例如综述, 临床数据, 文献索引, 参考文献等"
                    "## 目的: 请将下列图片所示的关于Ustekinumab(药物名)的医学报告翻译为专业的科学日文, 并按照下列规则排版输出"
                    "## 规则:"
                    "- 输出的一定是日文版本"
                    "- 按照原始图片中的排版生成Word"
                    "- 按照原文进行语句通顺的翻译, 表格里的内容也需要翻译为日文"
                    "- 标题左对齐显示, 药物名用原文, 主题需要翻译为日文"
                    "- 将输出的Word分为四个段落, 四个段落标题分别为'要約, 臨床データ, 文献検索, 参考', 要約是原文中的SUMMARY或COMPANY CORE DATA SHEET段落, 臨床データ为原文中的CLINICAL DATA或TECHNICAL MEMORANDUM段落, 文献検索是原文中的LITERATURE SEARCH段落, 参考是原文中的REFERENCES段落"
                    "- 若没有文献検索内容, 则输出'（このセクションは画像に含まれていません)'"
                    "- 若没有参考内容, 则输出'（このセクションは画像に含まれていません)'"
                    "- 文献检索和参考文献的内容不需要翻译,使用原文"
                    "- 段落标题需要翻译为日文"
                    "- 段落标题左对齐"
                    "- 删除掉页脚"
                    "- '•'是一个ListItem起始的标志, 'o'是属于上面ListItem的ListItem"
                    "- 如果有表格,则在输出表格最开始和最后加上'TABLE'几个字, 不要忘记表格描述和略语"
                    "- 如果没有表格,不输出'TABLE'"
                    "- 表格内容换行符用'-'表示"
                    "- 保留原文中的'•'符号"
                    "- 英文:Guttate Psoriasis 日文:滴状乾癬"
                    "- Word Format:"
                    "-- <1> 表格的格式和原始图片中的表格保持一致, 内容转换为日文"
                },
            ]
        },
        # {
        #     "role": "assistant",
        #     "content": "'The table has a total of"
        # }
    ],
    "max_tokens": 4080,
    "stream": False,
    "temperature": 0,
}

# Import PDF Image
for image_data in sImageData_list:
    url_data = {"type": "image_url", 
                "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
    dData["messages"][1]["content"].append(url_data)

# Make the API request
response = requests.post(
    f'{sEndpoint}openai/deployments/{sDeployment}/chat/completions?api-version=2024-02-01',
    headers={'api-key': sKey, 'Content-Type': 'application/json'},
    json=dData
)

# Print the response
print(response.json()['choices'][0]['message']['content'])

# COMMAND ----------

# MAGIC %md
# MAGIC HTML-Str_2_HTML

# COMMAND ----------

html_string = response.json()['choices'][0]['message']['content'].split("```html")[1].split("```")[0]
html_path = '/Workspace/Users/jfeng51@its.jnj.com/Transformed/HTML/' + file_name.split('.')[0] + "/"
if not os.path.exists(html_path):
    os.makedirs(html_path)
with open(html_path + file_name.split(".")[0] + '.html', 'w', encoding='utf-8') as f:
    f.write(html_string)

# COMMAND ----------

# MAGIC %md
# MAGIC 解析Word并填充

# COMMAND ----------

import re
from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

# 找表格||之间的'\n', 即找table行数
def count_newlines_between_pipes(str):
    first_pipe = str.find('|')
    last_pipe = str.rfind('|')
    if first_pipe == -1 or last_pipe == -1:
        return 0
    count = str[first_pipe:last_pipe].count('\n')
    return count

# 找表格|\n|之间的'|', 即找table列数
def count_segement_between_pipes(str):
    first_pipe = str.find("|")
    last_pipe = str.find("|\n")
    if first_pipe == -1 or last_pipe == -1:
        return 0
    count = str[first_pipe:last_pipe].count('|')
    return count

#设置单元格边框
def set_cell_border(cell: _Cell, **kwargs):
    """
    设置单元格边框

    参数：
    - cell (_Cell)：要设置边框的单元格
    - kwargs：边框属性，例如 top、bottom、left、right 等

    用法：
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#ff0000", "space": "0"},
        bottom={"sz": 12, "color": "#00ff00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 检查是否存在标签，如果不存在，则创建一个
    tcBorders = tcPr.first_child_found_in("w:tcborders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcborders')
        tcPr.append(tcBorders)

    # 遍历所有可用的标签
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # 检查是否存在标签，如果不存在，则创建一个
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # 设置属性
            for key in ("sz", "val", "color", "space", "shadow"):
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))    

# 正则表达判断(是否为字符)
def is_all_letters_regex(string):
    pattern = re.compile('^[a-zA-Z0-9]+$')
    return bool(pattern.match(string))

# 正则表达式判断(是否为非字符)
def is_only_non_letters_regex(string):
    pattern = re.compile('^[^a-zA-Z0-9]+$')
    return bool(pattern.match(string))

# 正则表达式判断(提取两个匹配的字符串中间的字符串)
def extract_between_strings(text, start_str, end_str):
    pattern = f"{re.escape(start_str)}(.*?){re.escape(end_str)}"
    match = re.search(pattern, text, re.DOTALL)
    if match:
        return match.group(1)
    else:
        return ""
    
# 正则表达式匹配'-'
def match_only_hyphens(s):
    pattern = r'^---$'
    if re.match(pattern, s):
        return True
    else:
        return False

def create_table(r, c):
    table = doc.add_table(rows=r, cols=c)
    return table

def parse_table(table, text):
    # table_text_list = text.split("|")
    table_text_list = text[text.find("|"): text.rfind("|")].strip().split("|")
    table_text_list = [item.strip() for item in table_text_list if item.strip() not in ["\n", '']]
    print(table_text_list)
    cell_cout = 0
    for row in range(len(table.rows)):    
        for col in range(len(table.columns)):
            cur_text = table_text_list[cell_cout]
            # is_empty = is_only_non_letters_regex(cur_text)
            is_empty = match_only_hyphens(cur_text)
            while is_empty:
                cell_cout +=1
                cur_text = table_text_list[cell_cout]
                # is_empty = is_only_non_letters_regex(cur_text)
                is_empty = match_only_hyphens(cur_text)
            table.cell(row, col).text = table_text_list[cell_cout].replace("-•", "\n•").strip()
            # print(table.cell(row, col).text)
            cell_cout += 1

def parse_data_no_table(main_table, summary_text, clinical_text): #func重载
    for row in main_table.rows:
        if row.cells[0].text == "QA番号":
            row.cells[1].text = "QA-0000001"
        elif row.cells[0].text == "領域/製品":
            row.cells[1].text = "免疫領域/ステラーラ/STL"
        elif row.cells[0].text == "QA種別":
            row.cells[1].text = "手動記入"
        elif row.cells[0].text == "質問":
            row.cells[1].text = "[STL]" + title_name
        elif row.cells[0].text == "回答\n（書式設定あり）":
            row.cells[1].paragraphs[1].text = summary_text.strip()
            row.cells[1].paragraphs[3].text = clinical_text.strip()
        elif row.cells[0].text == "引用文献":
            row.cells[1].text = "文献検索:\n" + literature_text + "\n" +"参考:\n" + reference_text
            row.cells[1].text = row.cells[1].text.strip()


def parse_data(main_table, new_table, summary_text, clinical_text): #func重载
    for row in main_table.rows:
        if row.cells[0].text == "QA番号":
            row.cells[1].text = "QA-0000001"
        elif row.cells[0].text == "領域/製品":
            row.cells[1].text = "免疫領域/ステラーラ/STL"
        elif row.cells[0].text == "QA種別":
            row.cells[1].text = "手動記入"
        elif row.cells[0].text == "質問":
            row.cells[1].text = "[STL]" + title_name
        elif row.cells[0].text == "回答\n（書式設定あり）":
            """ summary """
            row.cells[1].paragraphs[1].text = summary_text

            """ clinical """
            # Split Clinical Data
            before_table_text = clinical_text.split("TABLE", 1)[0].strip()
            # print(before_table_text)
            after_table_text = clinical_text.split("TABLE", 2)[2].strip()
            # print(after_table_text)
            table_describle_text = clinical_text[clinical_text.find('TABLE'):clinical_text.find('|')].strip()
            # print(table_describle_text)
            abbreviation_text = clinical_text[clinical_text.rfind('|'):clinical_text.rfind('TABLE')].strip()
            # print(abbreviation_text)

            #Table前
            row.cells[1].paragraphs[3].text = before_table_text

            #Table描述
            row.cells[1].add_paragraph()
            row.cells[1].paragraphs[4].text = table_describle_text

            #Table
            row.cells[1].add_paragraph()
            row.cells[1].paragraphs[5]._p.addnext(new_table._tbl)

            #Talbe略语
            row.cells[1].add_paragraph()
            row.cells[1].paragraphs[6].text = abbreviation_text

            #Table后
            row.cells[1].add_paragraph()
            row.cells[1].paragraphs[7].text = after_table_text
        elif row.cells[0].text == "引用文献":
            row.cells[1].text = "文献検索:\n" + literature_text + "\n" + "参考:\n" + reference_text
            row.cells[1].text = row.cells[1].text.strip()

# 指定 Word 文档路径
doc_path = '/Workspace/Users/jfeng51@its.jnj.com/Transformed/Template/Japanese_Template_1.docx'
save_path = '/Workspace/Users/jfeng51@its.jnj.com/Transformed/Word/' + file_name.split('.')[0] + '/'
if not os.path.exists(save_path):
    os.makedirs(save_path)
save_path = save_path + file_name.split('.')[0] + '.docx'

# Parse Data
text = response.json()['choices'][0]['message']['content'].split("```markdown")[1].split("```")[0].strip()
# print(text)

drag_name = text.split("\n")[0].strip()
drag_name = drag_name.replace("#", "").strip()
# print(drag_name + "\n")

title_name = text.split("\n", 2)[1].strip()
title_name = title_name.replace('#', "").strip()
# print(title_name + "\n")

summary_text = extract_between_strings(text, "要約", "臨床データ")
summary_text = summary_text.replace("#", "").strip()
# print(summary_text + "\n")

clinical_text = extract_between_strings(text, "臨床データ", "文献検索")
clinical_text = clinical_text.replace("#", "").strip()
# print(clinical_text + "\n")

literature_text = extract_between_strings(text, "文献検索", "参考")
literature_text = literature_text.replace("#", "").strip()
# print(literature_text + "\n")

reference_text = text.split("参考", 1)[1].strip()
# print(reference_text + "\n")

doc = Document(doc_path)
main_table = doc.tables[0]

if clinical_text.find("TABLE") != -1:
    table_text = clinical_text.split("TABLE", 2)[1].strip()
    # print(table_text)

    # Template Table
    main_table = doc.tables[0]
    target_cell = main_table.cell(4,1)
    # print(target_cell.text)

    # Create New Table
    new_table_row = count_newlines_between_pipes(table_text.strip())
    # print(new_table_row)
    new_table_col = count_segement_between_pipes(table_text.strip())
    # print(new_table_col)
    new_table = create_table(new_table_row, new_table_col)
    new_table.alignment = WD_TABLE_ALIGNMENT.CENTER  #设置新表格的对齐方式

    # Set New Table Boarder
    for row in new_table.rows:
        for cell in row.cells:
            set_cell_border(cell, 
                            top={"sz": 2, "val": "single", "color": "000000"}, 
                            bottom={"sz": 2, "val": "single", "color": "000000"},
                            left={"sz": 2, "val": "single", "color": "000000"},
                            right={"sz": 2, "val": "single", "color": "000000"})

    # Set New Table Content
    parse_table(new_table, table_text)

    # Fill Content
    parse_data(main_table, new_table, summary_text, clinical_text)
else: # No Table
    # Fill Content
    parse_data_no_table(main_table, summary_text, clinical_text)

# Save
doc.save(save_path)


# COMMAND ----------

# MAGIC %md
# MAGIC 提取HTML

# COMMAND ----------

""" Exract HTML 元素 """
# print(response.json()['choices'][0]['message']['content'].split("```html")[1].split("```")[0].strip())
html_str = response.json()['choices'][0]['message']['content'].split("```html")[1].split("```")[0].strip()
html_str = "<!DOCTYPE html>\n<html lang='ja'>\n<head>\n<meta charset='UTF-8'>\n</head>" + html_str
print(html_str)
with open('/Workspace/Users/jfeng51@its.jnj.com/Transformed/HTML/output_fr_str.html', 'w', encoding='utf-8') as f:
    f.write(html_str)



""" HTML转IMAGE """
# import imgkit
# imgkit.from_string(html_str, '/Workspace/Users/jfeng51@its.jnj.com/Transformed/output_image.jpg')


import imgkit

# 将 HTML 保存为图像
imgkit.from_file('/Workspace/Users/jfeng51@its.jnj.com/Transformed/HTML/output_fr_str.html',
                 '/Workspace/Users/jfeng51@its.jnj.com/Transformed/example.png',
                 options={'encoding': 'utf-8'})

# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------

from typing import Any
from pydantic import BaseModel
from unstructured.partition.pdf import partition_pdf
import google.protobuf

path = "/Workspace/Users/jfeng51@its.jnj.com/Transformed/"
file_name = "1482067.pdf"


"""  Text Table Image """
raw_pdf_elements = partition_pdf(
    filename=path + "PDF/" + file_name,                  # mandatory
    strategy="hi_res",                                     # mandatory to use ``hi_res`` strategy
    extract_images_in_pdf=True,                            # mandatory to set as ``True``
    infer_table_structure=True,
    # extract_image_block_types=["Image", "Table", "Text"],          # optional
    extract_image_block_types=["Image"],          # optional
    extract_image_block_output_dir=path + "images/" + file_name.split(".")[0] + "/",
    # extract_image_block_to_payload=False,                  # optional
    # chunking_strategy="by_title",
    # multipage_sections = "False",
    # include_page_breaks = "False",
)

""" Type """
# for ele in raw_pdf_elements:
#     print(ele.category)

""" Text """
# path = "/Workspace/Users/jfeng51@its.jnj.com/Transformed/"
# file_name = "885383.pdf"
# raw_pdf_elements = partition_pdf(path + "PDF/" + file_name,
#                                  strategy="auto")

# COMMAND ----------

image_list = []
for root, dirs, files in os.walk(path + "images/" + file_name.split(".")[0] + "/"):
    for file in files:
        file_path = os.path.join(root, file)
        image_list.append(file_path)

# COMMAND ----------

# TODO
# 按顺序将, image,table之前,之后以及相关的text保存并填充
from docx.shared import Inches
import re
from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

# 返回字符串位置
def find_str_positions(str, parse_str):
    positions = []
    index = str.find(parse_str)
    while index!= -1:
        positions.append(index)
        index = str.find(parse_str, index + 1)
    return positions

# 找表格||之间的'\n', 即找table行数
def count_newlines_between_pipes(str):
    first_pipe = str.find('|')
    last_pipe = str.rfind('|')
    if first_pipe == -1 or last_pipe == -1:
        return 0
    count = str[first_pipe:last_pipe].count('\n')
    return count

# 找表格|\n|之间的'|', 即找table列数
def count_segement_between_pipes(str):
    first_pipe = str.find("|")
    last_pipe = str.find("|\n")
    if first_pipe == -1 or last_pipe == -1:
        return 0
    count = str[first_pipe:last_pipe].count('|')
    return count

# 正则表达判断(是否为字符)
def is_all_letters_regex(string):
    pattern = re.compile('^[a-zA-Z0-9]+$')
    return bool(pattern.match(string))

# 正则表达式判断(是否为非字符)
def is_only_non_letters_regex(string):
    pattern = re.compile('^[^a-zA-Z0-9]+$')
    return bool(pattern.match(string))

# 正则表达式判断(提取两个匹配的字符串中间的字符串)
def extract_between_strings(text, start_str, end_str):
    pattern = f"{re.escape(start_str)}(.*?){re.escape(end_str)}"
    match = re.search(pattern, text, re.DOTALL)
    if match:
        return match.group(1)
    else:
        return ""
    
# 正则表达式匹配'-'
def match_only_hyphens(s):
    pattern = r'^---$'
    if re.match(pattern, s):
        return True
    else:
        return False

# 设置单元格边框
def set_cell_border(cell: _Cell, **kwargs):
    """
    设置单元格边框

    参数：
    - cell (_Cell)：要设置边框的单元格
    - kwargs：边框属性，例如 top、bottom、left、right 等

    用法：
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#ff0000", "space": "0"},
        bottom={"sz": 12, "color": "#00ff00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 检查是否存在标签，如果不存在，则创建一个
    tcBorders = tcPr.first_child_found_in("w:tcborders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcborders')
        tcPr.append(tcBorders)

    # 遍历所有可用的标签
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # 检查是否存在标签，如果不存在，则创建一个
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # 设置属性
            for key in ("sz", "val", "color", "space", "shadow"):
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))   

def create_table(r, c):
    table = doc.add_table(rows=r, cols=c)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  #设置新表格的对齐方式

    # Set New Table Boarder
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 
                            top={"sz": 2, "val": "single", "color": "000000"}, 
                            bottom={"sz": 2, "val": "single", "color": "000000"},
                            left={"sz": 2, "val": "single", "color": "000000"},
                            right={"sz": 2, "val": "single", "color": "000000"})
    return table 

# COMMAND ----------

# TODO
# 4o会将Image和Table搞混, 需要一个办法强制输出正确格式, 或删除Image相关的Prompt, 处理4o自然输出

import base64
import requests
from mimetypes import guess_type
from pdf2image import convert_from_path
from docx.shared import Inches
import os
 
# file_name = "Ustekinumab - Occurrence of Alopecia in Adult Patients with CD or UC Treated with Ustekinumab.pdf"
# file_name = "Ustekinumab - The Use of Vaccines in Infants Exposed In Utero to Ustekinumab.pdf"
# file_name = "Ustekinumab - Use in Adult Patients with Comorbid Liver Cirrhosis.pdf"
# file_name = "Ustekinumab - Use in Adult Patients with Crohn's Disease or Ulcerative Colitis and Comorbid PSC.pdf"
# file_name = "Ustekinumab - Occurrence of Cytomegalovirus in Adult Patients Receiving Ustekinumab.pdf"
file_name = "Ustekinumab - Occurrence of Skin-Related Adverse Events in Crohn's Disease and Ulcerative Colitis.pdf"

image_path = "/Workspace/Users/jfeng51@its.jnj.com/Transformed/Global_Image/" + file_name.split('.')[0] + "/"
pdf_path = "/Workspace/Users/jfeng51@its.jnj.com/Transformed/PDF/"

pdf_images_list = []
 
# 将PDF文件的每一页转换为图像
images = convert_from_path(pdf_path + file_name)

if not os.path.exists(image_path):
    os.makedirs(image_path)
 
# 可以遍历images列表来保存或处理每一张图像
for i, image in enumerate(images):
    # 图像保存路径
    image_name = f'page_{i + 1}.png'
    # 保存图像
    image.save(image_path + image_name, 'PNG')
    pdf_images_list.append(image_path + image_name)


# Encode the image to base64
sImageData_list = []
for image_path in pdf_images_list:
    sImageData = base64.b64encode(open(image_path, 'rb').read()).decode('utf-8')
    sImageData_list.append(sImageData)

# API Initialize
sEndpoint='https://azr-hjg-apac-ds-cae-openai-3.openai.azure.com/'
sKey='d3460cd8602e4a8d8420ba021b886cd3'
sDeployment='gpt-4o'

# dData is copied from the tutorial, Temelate
dData = {
    "messages": [
        {
            "role": "system",
            "content": "You are a professional Japanese translators for STL related treatments, your translated article is very fluent and easy to understand"
        },
        {
            "role": "user",
            "content": [
                {   
                    "type": "text",

                    "text": 
                    "## 角色: 专业医疗知识翻译员"
                    "## 背景: 下列图片是一份关于Ustekinumab药物的医疗报告, 其中包含一个标题(Ustekinumab + 主题), 以及一些段落标题例如综述, 临床数据, 文献索引, 参考文献等"
                    "## 目的: 请将下列图片所示的关于Ustekinumab(药物名)的医学报告翻译为专业的科学日文, 按照下列规则排版翻译输出"

                    "## 语言规则:"
                    "- 输出的一定是日文版本"
                    "- 表格里的内容需要翻译为日文"
                    "- TABLE内的内容需要翻译为日文"
                    "- 所有内容都需要翻译并输出为日文"
                    "- 按照原文进行语句通顺的翻译"
                    "- 翻译的日文尽量避免敬语, 用客观的陈述语句去翻译"
                    "- 译文精简"
                    "- 删除没有必要的重复的主语"
                    
                    "## 排版规则:"
                    "- 按照原始图片中的排版生成Word"
                    "- 输出格式为, ```markdown 换行 输出内容 换行 ```"
                    "- 主标题左对齐显示, 格式为药物名称 换行 标题， 药物名称用原文, 主题需要翻译"
                    "- 将输出的Word分为四个段落, 四个段落标题分别为'要約, 臨床データ, 文献検索, 参考', 要約是原文中的SUMMARY或COMPANY CORE DATA SHEET段落, 臨床データ为原文中的CLINICAL DATA或TECHNICAL MEMORANDUM段落, 文献検索是原文中的LITERATURE SEARCH段落, 参考是原文中的REFERENCES段落"
                    "- 若没有文献検索内容, 则输出'（このセクションは画像に含まれていません)'"
                    "- 若没有参考内容, 则输出'（このセクションは画像に含まれていません)'"
                    "- 段落标题需要翻译为日文"
                    "- 段落标题左对齐"
                    "- 删除掉页脚"
                    "- 输出保留原文中的'•'和'o'符号"
                    "- 如果有表格,则在输出表格最开始和最后加上'TABLE'几个字, 不要忘记表格描述和略语, 格式为: 'TABLE' 换行 表格描述 换行 表格(重要:表格内的内容也要翻译为日文) 换行 略语(重要:全部的Abbreviations内容) 'TABLE'(重要:结尾的'TABLE'一定要输出), Cell用'|'分割, 没有内容的Cell用|---|显示, 空的内容的Cell用|empty|表示, Cell内的文字需要翻译为日文并输出"
                    "- 如果没有表格,不输出'TABLE'"
                    "- 表格的格式和原始图片中的表格保持一致, 内容转换为日文"
                    "- 表格的Cell用|间隔"
                    "- 表格里的内容换行符用'-'表示"
                    # "- 如果有图片,则在输出图片最开始和最后加上'IMAGE'几个字, 不要忘记图片描述和略语, 格式为 IMAGE 换行 図:图片描述 换行 略语: 原文Abbreviations: 换行 IMAGE"
                    # "- 如果没有图片,不输出'IMAGE'"

                    "## 词汇规则:"
                    "- 英文:Guttate Psoriasis 日文:乾癬"
                    "- 英文:chronic plaque PsO 日文:尋常性乾癬"
                    "- 英文:Ustekinumab 日文:ステラーラ"
                },
            ]
        },
        # {
        #     "role": "assistant",
        #     "content": "'The table has a total of"
        # }
    ],
    "max_tokens": 4096,
    "stream": False,
    "temperature": 0,
}

# Import PDF Image
for image_data in sImageData_list:
    url_data = {"type": "image_url", 
                "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
    dData["messages"][1]["content"].append(url_data)

# Make the API request
response = requests.post(
    f'{sEndpoint}openai/deployments/{sDeployment}/chat/completions?api-version=2024-02-01',
    headers={'api-key': sKey, 'Content-Type': 'application/json'},
    json=dData
)

# Print the response
print(response.json()['choices'][0]['message']['content'])

# COMMAND ----------

"""
Multi Table
"""

def check_string_regex(s):
    pattern = r'^[|-]+$'
    if re.match(pattern, s):
        return True
    else:
        return False

# 找到所有的 key_word 出现的位置, 一一提取keyword之间字符串
def extract_keyword_strings(text, key_word):
    extract_contet = []
    table_indices = [m.start() for m in re.finditer(key_word, text)]
    if len(table_indices) < 2:
        return "没有足够的 TABLE 来截取"
    for key_word_index in range(len(table_indices) - 1):
        star_index = table_indices[key_word_index] + len(key_word)
        end_index = table_indices[key_word_index + 1]
        extract_contet.append(text[star_index:end_index].strip())
    return extract_contet

def parse_table(clinical_text, cell):
    clinical_text_list = clinical_text.split("TABLE")
    del(clinical_text_list[0])
    del(clinical_text_list[-1])
    table_list = extract_keyword_strings(clinical_text, "TABLE")
    for table_index in range(len(table_list)):
        if table_index % 2 == 0:
            """ 获取'TABLE'间内容 """
            table_text = table_list[table_index]

            """ TABLE描述和略语 """
            table_describle_text = table_text.split('|')[0].strip()
            # print(table_describle_text)
            abbreviation_text = table_text[table_text.rfind('|'):].replace("|", "").strip()
            # print(abbreviation_text)

            """ CREATE TABLE """
            # Create New Table
            new_table_row = count_newlines_between_pipes(table_text.strip())
            print("new_talbe_row: " + str(new_table_row))
            new_table_col = count_segement_between_pipes(table_text.strip())
            print("new_talbe_col: " + str(new_table_col))
            new_table = create_table(new_table_row, new_table_col)
            new_table.alignment = WD_TABLE_ALIGNMENT.CENTER  #设置新表格的对齐方式

            """ 分解TABLE内容(易出错) """
            # 方案一 '|'切割
            # table_text_list = table_text[table_text.find("|"): table_text.rfind("|")].strip().split("|")
            # table_text_list = [item.strip() for item in table_text_list if item.strip() not in ["\n", '']]

            # 方案二 '\n'切割
            table_content = (table_text[table_text.find("|"): table_text.rfind("|")] + '|').strip()
            row_list_content = table_content.split("\n")
            row_list_content = [item for item in row_list_content if not check_string_regex(item)]
            # print(row_list_content)

            """ 根据原有TABLE内容填充到NewTable """
            for row in range(len(new_table.rows)):
                content_list = []
                if row_list_content[row].count('|') != new_table_col + 1:
                    content_list = row_list_content[row].split('|')
                    del(content_list[0])
                    del(content_list[-1])
                    for i in range(new_table_col - len(content_list)):
                        content_list.append("")
                    print(content_list)
                else:
                    content_list = row_list_content[row].split('|')
                    del(content_list[0])
                    del(content_list[-1])
                    print(content_list)
                for col in range(len(new_table.columns)):
                    cur_text = content_list[col]
                    new_table.cell(row, col).text = content_list[col].replace("-•", "\n•").strip()
                    print(new_table.cell(row, col).text)
                    
            """ 合并段落 """
            cell.add_paragraph(table_describle_text)
            paragraph = cell.add_paragraph()
            paragraph._p.addnext(new_table._tbl)
            cell.add_paragraph(abbreviation_text)
        else:
            cell.add_paragraph(table_list[table_index].strip())
            # print(table_list[table_index])

def parse_data_no_table(main_table, summary_text, clinical_text, literature_text, reference_text):
    for row in main_table.rows:
        if row.cells[0].text == "QA番号":
            row.cells[1].text = "QA-0000001"
        elif row.cells[0].text == "領域/製品":
            row.cells[1].text = "免疫領域/ステラーラ/STL"
        elif row.cells[0].text == "QA種別":
            row.cells[1].text = "手動記入"
        elif row.cells[0].text == "質問":
            row.cells[1].text = "[STL]" + title_name
        elif row.cells[0].text == "回答\n（書式設定あり）":
            row.cells[1].paragraphs[2].text = summary_text.strip()
            row.cells[1].add_paragraph(clinical_text.strip())
            # row.cells[1].paragraphs[4].text = clinical_text.strip()
        elif row.cells[0].text == "引用文献":
            row.cells[1].text = "文献検索:\n" + literature_text + "\n" +"参考:\n" + reference_text
            row.cells[1].text = row.cells[1].text.strip()


def parse_data(main_table, summary_text, clinical_text, literature_text, reference_text):
    for row in main_table.rows:
        if row.cells[0].text == "QA番号":
            row.cells[1].text = "QA-0000001"
        elif row.cells[0].text == "領域/製品":
            row.cells[1].text = "免疫領域/ステラーラ/STL"
        elif row.cells[0].text == "QA種別":
            row.cells[1].text = "手動記入"
        elif row.cells[0].text == "質問":
            row.cells[1].text = "[STL]" + title_name
        elif row.cells[0].text == "回答\n（書式設定あり）":
            """ summary """
            row.cells[1].paragraphs[2].text = summary_text.strip()

            """ clinical """
            before_table_text = clinical_text[:clinical_text.find("TABLE")].strip()
            after_table_text = clinical_text[clinical_text.rfind("TABLE"):].replace("TABLE", "").strip()
            
            # print(before_table_text)
            # print(after_table_text)

            row.cells[1].add_paragraph(before_table_text)
            parse_table(clinical_text, row.cells[1])
            row.cells[1].add_paragraph(after_table_text)

        elif row.cells[0].text == "引用文献":
            row.cells[1].text = "文献検索:\n" + literature_text + "\n" + "参考:\n" + reference_text
            row.cells[1].text = row.cells[1].text.strip()

# 指定 Word 文档路径
doc_path = '/Workspace/Users/jfeng51@its.jnj.com/Transformed/Template/Japanese_Template_1.docx'
save_path = '/Workspace/Users/jfeng51@its.jnj.com/Transformed/Word/' + file_name.split('.')[0] + '/'
if not os.path.exists(save_path):
    os.makedirs(save_path)
save_path = save_path + file_name.split('.')[0] + '.docx'

# Parse Data
text = response.json()['choices'][0]['message']['content'].split("```markdown")[1].split("```")[0].strip()
# print(text)

drag_name = text.split("\n")[0].strip()
drag_name = drag_name.replace("#", "").strip()
# print(drag_name + "\n")

title_name = text.split("\n", 2)[1].strip()
title_name = title_name.replace('#', "").strip()
# print(title_name + "\n")

summary_text = extract_between_strings(text, "要約", "臨床データ")
summary_text = summary_text.replace("#", "").strip()
# print(summary_text + "\n")

clinical_text = extract_between_strings(text, "臨床データ", "文献検索")
clinical_text = clinical_text.replace("#", "").strip()
# print(clinical_text + "\n")

literature_text = extract_between_strings(text, "文献検索", "参考")
literature_text = literature_text.replace("#", "").strip()
# print(literature_text + "\n")

reference_text = text.split("参考", 1)[1].strip()
# print(reference_text + "\n")

doc = Document(doc_path)
main_table = doc.tables[0]
target_cell = main_table.cell(4,1)

if clinical_text.find("TABLE") != -1:
    parse_data(main_table, summary_text, clinical_text, literature_text, reference_text)
else:
    parse_data_no_table(main_table, summary_text, clinical_text, literature_text, reference_text)

# Save
doc.save(save_path)

# COMMAND ----------

"""
MULTI IMAGE
"""

# TODO
# 按顺序将, image,table之前,之后以及相关的text保存并填充
from docx.shared import Inches
import re
from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

def parse_data(main_table, summary_text, clinical_text):
    for row in main_table.rows:
        if row.cells[0].text == "QA番号":
            row.cells[1].text = "QA-0000001"
        elif row.cells[0].text == "領域/製品":
            row.cells[1].text = "免疫領域/ステラーラ/STL"
        elif row.cells[0].text == "QA種別":
            row.cells[1].text = "手動記入"
        elif row.cells[0].text == "質問":
            row.cells[1].text = "[STL]" + title_name
        elif row.cells[0].text == "回答\n（書式設定あり）":
            """ summary """
            row.cells[1].paragraphs[1].text = summary_text

            """ clinical """
            paragraphs_index=3
            # Split Clinical Data
            non_image_text1 = clinical_text.split("IMAGE", 4)[0].strip()
            non_image_text2 = clinical_text.split("IMAGE", 4)[2].strip()
            non_image_text3 = clinical_text.split("IMAGE", 4)[4].strip()

            image_text_1 = clinical_text.split("IMAGE", 4)[1].strip()
            image_text_2 = clinical_text.split("IMAGE", 4)[3].strip()

            row.cells[1].paragraphs[paragraphs_index].text = non_image_text1

            row.cells[1].add_paragraph()
            paragraphs_index += 1
            image1_describe = image_text_1[image_text_1.find("図:") : image_text_1.find("略語:")].replace("図:", "").strip()
            image1_abbreviations = image_text_1[image_text_1.find("略語:") : -1].replace("略語:", "").strip()

            row.cells[1].paragraphs[paragraphs_index].text = "\n" + image1_describe

            row.cells[1].add_paragraph()
            paragraphs_index += 1
            run_1_2 = row.cells[1].paragraphs[paragraphs_index].add_run()
            run_1_2.add_picture('/Workspace/Users/jfeng51@its.jnj.com/Transformed/images/1482067/figure-2-1.jpg', width=Inches(6))

            row.cells[1].add_paragraph()
            paragraphs_index += 1
            row.cells[1].paragraphs[paragraphs_index].text = image1_abbreviations

            row.cells[1].add_paragraph()
            paragraphs_index += 1
            row.cells[1].paragraphs[paragraphs_index].text = non_image_text2

            row.cells[1].add_paragraph()
            paragraphs_index += 1
            image2_describe = image_text_2[image_text_2.find("図:") : image_text_2.find("略語:")].replace("図:", "").strip()
            image2_abbreviations = image_text_2[image_text_2.find("略語:") : -1].replace("略語:", "").strip()

            row.cells[1].paragraphs[paragraphs_index].text = "\n" + image2_describe

            row.cells[1].add_paragraph()
            paragraphs_index += 1
            run_2_2 = row.cells[1].paragraphs[paragraphs_index].add_run()
            run_2_2.add_picture('/Workspace/Users/jfeng51@its.jnj.com/Transformed/images/1482067/figure-3-2.jpg', width=Inches(6))

            row.cells[1].add_paragraph()
            paragraphs_index += 1
            row.cells[1].paragraphs[paragraphs_index].text = image2_abbreviations

            row.cells[1].add_paragraph()
            paragraphs_index += 1
            row.cells[1].paragraphs[paragraphs_index].text = non_image_text3

        elif row.cells[0].text == "引用文献":
            row.cells[1].text = "文献検索:\n" + literature_text + "\n" + "参考:\n" + reference_text
            row.cells[1].text = row.cells[1].text.strip()

# 指定 Word 文档路径
doc_path = '/Workspace/Users/jfeng51@its.jnj.com/Transformed/Template/Japanese_Template_1.docx'
save_path = '/Workspace/Users/jfeng51@its.jnj.com/Transformed/Word/' + file_name.split('.')[0] + '/'
if not os.path.exists(save_path):
    os.makedirs(save_path)
save_path = save_path + file_name.split('.')[0] + '.docx'

# Parse Data
text = response.json()['choices'][0]['message']['content'].split("```markdown")[1].split("```")[0].strip()
# print(text)

drag_name = text.split("\n")[0].strip()
drag_name = drag_name.replace("#", "").strip()
# print(drag_name + "\n")

title_name = text.split("\n", 2)[1].strip()
title_name = title_name.replace('#', "").strip()
# print(title_name + "\n")

summary_text = extract_between_strings(text, "要約", "臨床データ")
summary_text = summary_text.replace("#", "").strip()
# print(summary_text + "\n")

clinical_text = extract_between_strings(text, "臨床データ", "文献検索")
clinical_text = clinical_text.replace("#", "").strip()
# print(clinical_text + "\n")

literature_text = extract_between_strings(text, "文献検索", "参考")
literature_text = literature_text.replace("#", "").strip()
# print(literature_text + "\n")

reference_text = text.split("参考", 1)[1].strip()
# print(reference_text + "\n")

doc = Document(doc_path)
main_table = doc.tables[0]

if clinical_text.find("IMAGE") != -1:
    parse_data(main_table, summary_text, clinical_text)

# Save
doc.save(save_path)

# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------



# COMMAND ----------


